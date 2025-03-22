import pyaudio
import wave
import threading
import tkinter as tk
from tkinter import messagebox
import speech_recognition as sr
import google.generativeai as genai
from docx import Document

# Configure Gemini API
genai.configure(api_key="AIzaSyAMgHHUs81sI6T3f6eJxK_fD6jtL4yML40")

# Global variables
recording = False
audio_filename = "meeting_audio.wav"

def start_recording():
    """Start recording audio in a separate thread"""
    global recording
    recording = True
    threading.Thread(target=record_audio).start()
    start_button.config(state=tk.DISABLED)
    stop_button.config(state=tk.NORMAL)

def stop_recording():
    """Stop recording audio"""
    global recording
    recording = False
    stop_button.config(state=tk.DISABLED)
    start_button.config(state=tk.NORMAL)
    messagebox.showinfo("Recording Stopped", "Audio saved! Processing transcript...")
    
    # Convert speech to text
    transcript = speech_to_text(audio_filename)
    
    # Enhance with Gemini AI
    if transcript:
        enhanced_text = enhance_text_with_gemini(transcript)
        save_transcript_to_word(enhanced_text)
        messagebox.showinfo("Transcription", f"Enhanced Transcript saved to 'meeting_transcript.docx'")
    else:
        messagebox.showerror("Error", "No speech detected!")

def record_audio():
    """Record audio until manually stopped"""
    global recording
    p = pyaudio.PyAudio()
    stream = p.open(format=pyaudio.paInt16, channels=1, rate=44100, input=True, frames_per_buffer=1024)
    frames = []

    print("🎙 Recording started... Press Stop to end.")
    while recording:
        frames.append(stream.read(1024))

    stream.stop_stream()
    stream.close()
    p.terminate()

    # Save the recorded audio
    wf = wave.open(audio_filename, 'wb')
    wf.setnchannels(1)
    wf.setsampwidth(p.get_sample_size(pyaudio.paInt16))
    wf.setframerate(44100)
    wf.writeframes(b''.join(frames))
    wf.close()
    
    print(f"✅ Audio saved as {audio_filename}")

def speech_to_text(filename):
    """Convert recorded audio to text"""
    recognizer = sr.Recognizer()
    with sr.AudioFile(filename) as source:
        print("🎧 Processing audio...")
        audio = recognizer.record(source)

    try:
        text = recognizer.recognize_google(audio, language="hi-IN,en-IN")
        print(f"📝 Recognized Text: {text}")
        return text
    except sr.UnknownValueError:
        print("❌ Could not understand audio.")
        return ""
    except sr.RequestError:
        print("❌ API unavailable.")
        return ""

def enhance_text_with_gemini(text):
    """Enhance transcript using Gemini AI"""
    model = genai.GenerativeModel("gemini-1.5-pro-latest")
    response = model.generate_content(f"Summarize and structure this meeting transcript:\n{text}")
    return response.text.strip()

def save_transcript_to_word(text, filename="meeting_transcript.docx"):
    """Save the transcript to a Word file"""
    doc = Document()
    doc.add_heading("Meeting Transcript", level=1)
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"📄 Transcript saved as {filename}")

# GUI using Tkinter
root = tk.Tk()
root.title("Meeting Recorder")
root.geometry("300x150")

start_button = tk.Button(root, text="Start Recording", command=start_recording, bg="green", fg="white")
stop_button = tk.Button(root, text="Stop Recording", command=stop_recording, bg="red", fg="white", state=tk.DISABLED)

start_button.pack(pady=10)
stop_button.pack(pady=10)

root.mainloop()